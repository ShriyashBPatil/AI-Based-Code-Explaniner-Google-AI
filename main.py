import google.generativeai as genai
import customtkinter as ctk
from tkinter import filedialog
import time
from docx import Document 
import threading  

def configure_genai(api_key):
    genai.configure(api_key=api_key)
    generation_config = {
        "temperature": 1,
        "top_p": 0.95,
        "top_k": 64,
        "max_output_tokens": 8192,
        "response_mime_type": "text/plain",
    }
    model = genai.GenerativeModel(
        model_name="gemini-1.5-flash",
        generation_config=generation_config,
    )
    return model


api_key = "AIzaSyA6DZkB7w66xqcUXNCuYiM6f51rTOevOKc"
model = configure_genai(api_key)

chat_history_list = []

def format_history(history):
    return [{"parts": [{"text": entry["content"]}], "role": entry["role"]} for entry in history]

def send_message(message, history): 
    formatted_history = format_history(history)
    chat_session = model.start_chat(
        history=formatted_history
    )
    response = chat_session.send_message(message)
    return response.text

def type_text(widget, text, delay=10):
    def update_text(index=0):
        if index < len(text):
            widget.configure(state=ctk.NORMAL)
            if text.startswith("**"):
                widget.insert(ctk.END, text[index], "bold")
            else:
                widget.insert(ctk.END, text[index])
            widget.see(ctk.END)
            widget.configure(state=ctk.DISABLED)
            widget.after(delay, update_text, index + 1)
        else:
            widget.insert(ctk.END, "\n")
            widget.configure(state=ctk.DISABLED)

    update_text()

# Function to open the About window
def open_about_window():
    about_window = ctk.CTkToplevel(root)
    about_window.title("About")
    about_window.geometry("400x200")
    about_window.attributes("-topmost", True)

    about_label = ctk.CTkLabel(about_window, text="About This App", font=("Arial", 16, "bold"))
    about_label.pack(pady=10)

    creator_info = (
        "Created by: Your Name\n"
        "Version: 1.0\n"
        "Date: October 2023\n"
        "Description: This app provides detailed explanations of your code."
    )
    info_label = ctk.CTkLabel(about_window, text=creator_info, font=("Arial", 12), justify="left")
    info_label.pack(pady=10, padx=10)

# Function to show error messages
def show_error_message(error_message):
    error_window = ctk.CTkToplevel(root)
    error_window.title("Error")
    error_window.geometry("300x150")

    error_label = ctk.CTkLabel(error_window, text="An error occurred:", font=("Arial", 14, "bold"))
    error_label.pack(pady=10)

    message_label = ctk.CTkLabel(error_window, text=error_message, font=("Arial", 12), justify="left", wraplength=250)
    message_label.pack(pady=5, padx=10)

# Function to handle sending messages
def on_send():
    print("Send button clicked")
    def send_message_thread():
        try:
            user_message = user_input.get("1.0", ctk.END).strip()
            print(f"User message: {user_message}")
            if user_message:
                chat_history.configure(state=ctk.NORMAL)
                chat_history.insert(ctk.END, "You: ", "bold")
                chat_history.insert(ctk.END, user_message + "\n")
                chat_history.configure(state=ctk.DISABLED)
                user_input.delete("1.0", ctk.END)

                chat_history_list.append({"role": "user", "content": user_message})

                bot_response = send_message(user_message, chat_history_list)
                print(f"Bot response: {bot_response}")

                chat_history_list.append({"role": "model", "content": bot_response})

                chat_history.configure(state=ctk.NORMAL)
                chat_history.insert(ctk.END, "AI: ", "bold")
                chat_history.configure(state=ctk.DISABLED)
                type_text(chat_history, bot_response)

                current_input.configure(state=ctk.NORMAL)
                current_input.insert(ctk.END, "You: ", "bold")
                current_input.insert(ctk.END, user_message + "\n")
                current_input.insert(ctk.END, "AI: ", "bold")
                current_input.insert(ctk.END, bot_response + "\n")
                current_input.configure(state=ctk.DISABLED)
        except Exception as e:
            print(f"Error in send_message_thread: {e}")
            show_error_message(str(e))

    threading.Thread(target=send_message_thread).start()

def export_to_docx(content):
    doc = Document()
    doc.add_heading('Code Explanation', 0)
    doc.add_paragraph(content)
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if file_path:
        doc.save(file_path)

def open_file():
    def open_file_thread():
        file_path = filedialog.askopenfilename()
        if file_path:
            with open(file_path, 'r') as file:
                file_content = file.read()
                code_display.configure(state=ctk.NORMAL)
                code_display.delete(1.0, ctk.END)
                code_display.insert(ctk.END, file_content)
                code_display.configure(state=ctk.DISABLED)
                
                explanation_request = f"Explain this code in detail:\n{file_content}"
                bot_response = send_message(explanation_request, chat_history_list)
           
                chat_history.configure(state=ctk.NORMAL)
                chat_history.insert(ctk.END, "Code:\n", "bold")
                chat_history.insert(ctk.END, file_content + "\n\n")
                chat_history.insert(ctk.END, "Bot: ", "bold")
                chat_history.configure(state=ctk.DISABLED)
                type_text(chat_history, bot_response)
              
                global last_explanation
                last_explanation = bot_response

            update_line_numbers()
    
    threading.Thread(target=open_file_thread).start()

def update_line_numbers(event=None):
    code_content = code_display.get("1.0", "end-1c")
    lines = code_content.split('\n')
    line_numbers = "\n".join(str(i + 1) for i in range(len(lines)))
    line_numbers_display.configure(state=ctk.NORMAL)
    line_numbers_display.delete("1.0", ctk.END)
    line_numbers_display.insert("1.0", line_numbers)
    line_numbers_display.configure(state=ctk.DISABLED)

def on_scroll(event):
    code_display.yview_scroll(int(-1*(event.delta/120)), "units")
    line_numbers_display.yview_scroll(int(-1*(event.delta/120)), "units")

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("green")

root = ctk.CTk()
root.title("Code Explainer")
root.geometry("1270x720")

title_label = ctk.CTkLabel(root, text="Code Explainer", font=("Arial", 24, "bold"), text_color="#2E8B57")
title_label.grid(row=0, column=0, columnspan=2, pady=10, sticky="ew")

description_label = ctk.CTkLabel(root, text="Get detailed explanations of your code", font=("Arial", 16), text_color="#555555")
description_label.grid(row=1, column=0, columnspan=2, pady=5, sticky="ew")

main_frame = ctk.CTkFrame(root)
main_frame.grid(row=2, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")

root.grid_rowconfigure(2, weight=1)
root.grid_columnconfigure(0, weight=1)
root.grid_columnconfigure(1, weight=1)

chat_history_label = ctk.CTkLabel(main_frame, text="Current Input", font=("Arial", 14, "bold"), text_color="#2E8B57")
chat_history_label.grid(row=0, column=0, pady=5, sticky="ew")
chat_history = ctk.CTkTextbox(main_frame, state='disabled', wrap=ctk.WORD, font=("Arial", 12), text_color="#333333")
chat_history.grid(row=1, column=0, padx=5, pady=5, sticky="nsew")

current_input_label = ctk.CTkLabel(main_frame, text="Chat History", font=("Arial", 14, "bold"), text_color="#2E8B57")
current_input_label.grid(row=2, column=0, pady=5, sticky="ew")
current_input = ctk.CTkTextbox(main_frame, state='disabled', wrap=ctk.WORD, font=("Arial", 12), text_color="#333333")
current_input.grid(row=3, column=0, padx=5, pady=5, sticky="nsew")

code_display_label = ctk.CTkLabel(main_frame, text="Code Display", font=("Arial", 14, "bold"), text_color="#2E8B57")
code_display_label.grid(row=0, column=1, pady=5, sticky="ew")

code_frame = ctk.CTkFrame(main_frame)
code_frame.grid(row=1, column=1, rowspan=2, padx=10, pady=10, sticky="nsew")

line_numbers_display = ctk.CTkTextbox(code_frame, width=4, state='disabled', wrap=ctk.NONE, font=("Arial", 10), text_color="#333333")
line_numbers_display.pack(side=ctk.LEFT, fill=ctk.Y)

code_display = ctk.CTkTextbox(code_frame, state='disabled', wrap=ctk.WORD, font=("Arial", 12), text_color="#333333")
code_display.pack(side=ctk.RIGHT, fill=ctk.BOTH, expand=True)

user_input_label = ctk.CTkLabel(main_frame, text="Your Input", font=("Arial", 14, "bold"), text_color="#2E8B57")
user_input_label.grid(row=2, column=1, pady=5, sticky="ew")

user_input = ctk.CTkTextbox(main_frame, font=("Arial", 12), text_color="#333333")
user_input.grid(row=3, column=1, padx=5, pady=5, sticky="nsew")

main_frame.grid_rowconfigure(1, weight=1)
main_frame.grid_rowconfigure(3, weight=1)
main_frame.grid_columnconfigure(0, weight=1)
main_frame.grid_columnconfigure(1, weight=1)

button_frame = ctk.CTkFrame(root)
button_frame.grid(row=3, column=0, columnspan=2, padx=5, pady=5, sticky="ew")

send_button = ctk.CTkButton(button_frame, text="Send", command=on_send, font=("Arial", 12, "bold"), text_color="#FFFFFF", fg_color="#4CAF50")
send_button.pack(side=ctk.LEFT, padx=5, pady=5, fill=ctk.X, expand=True)

file_button = ctk.CTkButton(button_frame, text="Upload File", command=open_file, font=("Arial", 12, "bold"), text_color="#FFFFFF", fg_color="#2196F3")
file_button.pack(side=ctk.LEFT, padx=5, pady=5, fill=ctk.X, expand=True)

export_button = ctk.CTkButton(button_frame, text="Export Explanation", command=lambda: export_to_docx(last_explanation), font=("Arial", 12, "bold"), text_color="#FFFFFF", fg_color="#FF9800")
export_button.pack(side=ctk.LEFT, padx=5, pady=5, fill=ctk.X, expand=True)

export_chat_button = ctk.CTkButton(button_frame, text="Export Chat History", command=lambda: export_chat_history_to_docx(chat_history_list), font=("Arial", 12, "bold"), text_color="#FFFFFF", fg_color="#FF9800")
export_chat_button.pack(side=ctk.LEFT, padx=5, pady=5, fill=ctk.X, expand=True)

about_button = ctk.CTkButton(button_frame, text="About", command=open_about_window, font=("Arial", 12, "bold"), text_color="#FFFFFF", fg_color="#4CAF50")
about_button.pack(side=ctk.LEFT, padx=5, pady=5, fill=ctk.X, expand=True)

def export_chat_history_to_docx(chat_history_list):
    doc = Document()
    doc.add_heading('Chat History', 0)
    for entry in chat_history_list:
        role = "You" if entry["role"] == "user" else "Bot"
        doc.add_paragraph(f"{role}: {entry['content']}")
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if file_path:
        doc.save(file_path)

root.mainloop()